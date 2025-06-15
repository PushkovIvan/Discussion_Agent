import streamlit as st
import json
import io
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from langchain_core.messages import HumanMessage, SystemMessage, AIMessage
from langchain_community.chat_models import GigaChat
from langchain.agents import AgentExecutor, create_openai_functions_agent
from langchain.tools import BaseTool
from langchain.schema import BaseMessage
from typing import List, Dict, Any, Optional
import re

from config import CONFIG

# Инициализация GigaChat
giga = GigaChat(
    scope='GIGACHAT_API_CORP',
    credentials=CONFIG["token"]["gigachat"],
    verify_ssl_certs=False,
    model="GigaChat-2-Max"
)

class SituationGenerator:
    """Агент для создания ситуаций и ролей"""
    
    def __init__(self, llm):
        self.llm = llm
    
    def generate_situation(self) -> Dict[str, str]:
        """Генерирует случайную ситуацию для переговоров"""
        system_prompt = """
        Ты эксперт по созданию реалистичных сценариев для тренировки навыков переговоров.
        Создай интересную и сложную ситуацию для переговоров между менеджером и клиентом.
        Ситуация должна быть реалистичной и требовать навыков убеждения и работы с возражениями.
        """
        
        user_prompt = """
        Создай ситуацию для тренировки переговоров. Верни ответ в формате JSON:
        {
            "situation": "описание ситуации",
            "manager_role": "роль менеджера",
            "client_role": "роль клиента",
            "manager_goal": "цель менеджера",
            "client_concerns": "основные возражения клиента",
            "product": "продукт или услуга",
            "context": "дополнительный контекст"
        }
        """
        
        messages = [
            SystemMessage(content=system_prompt),
            HumanMessage(content=user_prompt)
        ]
        
        response = self.llm.invoke(messages)
        
        try:
            # Извлекаем JSON из ответа
            json_match = re.search(r'\{.*\}', response.content, re.DOTALL)
            if json_match:
                return json.loads(json_match.group())
            else:
                # Fallback если JSON не найден
                return self._create_default_situation()
        except:
            return self._create_default_situation()
    
    def _create_default_situation(self) -> Dict[str, str]:
        return {
            "situation": "Переговоры по внедрению CRM-системы в среднюю компанию",
            "manager_role": "Менеджер по продажам IT-решений",
            "client_role": "Директор по развитию компании",
            "manager_goal": "Продать CRM-систему стоимостью 500,000 рублей",
            "client_concerns": "Высокая стоимость, сложность внедрения, сомнения в ROI",
            "product": "CRM-система для автоматизации продаж",
            "context": "Клиент уже использует простую систему учета, но хочет масштабировать бизнес"
        }

class DialogueAgent:
    """Агент для ведения диалога от имени клиента"""
    
    def __init__(self, llm, situation: Dict[str, str]):
        self.llm = llm
        self.situation = situation
    
    def respond_as_client(self, conversation_history: List[Dict[str, str]]) -> str:
        """Отвечает от имени клиента на основе истории диалога"""
        system_prompt = f"""
        Ты играешь роль клиента в переговорах. Твоя роль: {self.situation['client_role']}.
        Ситуация: {self.situation['situation']}
        Твои основные возражения: {self.situation['client_concerns']}
        
        Веди себя как реальный клиент:
        - Задавай сложные вопросы
        - Выражай сомнения и возражения
        - Будь скептичен к обещаниям
        - Требуй конкретные доказательства
        - Не соглашайся легко
        - Отвечай кратко и по делу
        """
        
        # Формируем историю диалога
        history_text = ""
        for msg in conversation_history:
            if msg['role'] == 'manager':
                history_text += f"Менеджер: {msg['content']}\n"
            elif msg['role'] == 'client':
                history_text += f"Клиент: {msg['content']}\n"
        
        user_prompt = f"""
        История диалога:
        {history_text}
        
        Ответь как клиент на последнее сообщение менеджера. Будь реалистичным и сложным собеседником.
        """
        
        messages = [
            SystemMessage(content=system_prompt),
            HumanMessage(content=user_prompt)
        ]
        
        response = self.llm.invoke(messages)
        return response.content.strip()

class DialogueEndDetector:
    """Агент для определения завершения диалога"""
    
    def __init__(self, llm, situation: Dict[str, str]):
        self.llm = llm
        self.situation = situation
    
    def should_end_dialogue(self, conversation_history: List[Dict[str, str]]) -> bool:
        """Определяет, пора ли завершить диалог"""
        if len(conversation_history) < 4:  # Минимум 4 сообщения
            return False
        
        system_prompt = f"""
        Ты эксперт по анализу переговоров. Определи, достигнута ли цель переговоров.
        Цель менеджера: {self.situation['manager_goal']}
        
        Диалог должен завершиться, если:
        1. Клиент согласился на предложение менеджера
        2. Клиент категорически отказался и дальнейшие переговоры бессмысленны
        3. Достигнут компромисс, удовлетворяющий обе стороны
        4. Диалог зашел в тупик и требует перерыва
        """
        
        history_text = ""
        for msg in conversation_history:
            if msg['role'] == 'manager':
                history_text += f"Менеджер: {msg['content']}\n"
            elif msg['role'] == 'client':
                history_text += f"Клиент: {msg['content']}\n"
        
        user_prompt = f"""
        История диалога:
        {history_text}
        
        Ответь только "ДА" если диалог пора завершать, или "НЕТ" если нужно продолжить.
        """
        
        messages = [
            SystemMessage(content=system_prompt),
            HumanMessage(content=user_prompt)
        ]
        
        response = self.llm.invoke(messages)
        return "ДА" in response.content.upper()

class ReportGenerator:
    """Агент для создания отчета по результатам переговоров"""
    
    def __init__(self, llm, situation: Dict[str, str]):
        self.llm = llm
        self.situation = situation
    
    def generate_report(self, conversation_history: List[Dict[str, str]]) -> Dict[str, str]:
        """Генерирует подробный отчет по переговорам"""
        system_prompt = f"""
        Ты эксперт по анализу переговоров. Создай подробный отчет по результатам диалога.
        Ситуация: {self.situation['situation']}
        Цель менеджера: {self.situation['manager_goal']}
        Продукт: {self.situation['product']}
        
        Проанализируй только сообщения менеджера и дай рекомендации по улучшению навыков.
        """
        
        # Извлекаем только сообщения менеджера
        manager_messages = [msg['content'] for msg in conversation_history if msg['role'] == 'manager']
        manager_text = "\n".join([f"Сообщение {i+1}: {msg}" for i, msg in enumerate(manager_messages)])
        
        user_prompt = f"""
        Сообщения менеджера:
        {manager_text}
        
        Создай отчет в формате JSON:
        {{
            "summary": "краткое резюме переговоров",
            "strengths": ["сильные стороны менеджера"],
            "weaknesses": ["слабые стороны и ошибки"],
            "recommendations": ["конкретные рекомендации по улучшению"],
            "growth_areas": ["области для развития"],
            "techniques_used": ["использованные техники"],
            "missed_opportunities": ["упущенные возможности"],
            "overall_rating": "общая оценка от 1 до 10"
        }}
        """
        
        messages = [
            SystemMessage(content=system_prompt),
            HumanMessage(content=user_prompt)
        ]
        
        response = self.llm.invoke(messages)
        
        try:
            json_match = re.search(r'\{.*\}', response.content, re.DOTALL)
            if json_match:
                return json.loads(json_match.group())
            else:
                return self._create_default_report()
        except:
            return self._create_default_report()
    
    def _create_default_report(self) -> Dict[str, str]:
        return {
            "summary": "Проведены переговоры по продаже продукта",
            "strengths": ["Активное слушание", "Профессиональный подход"],
            "weaknesses": ["Недостаточно аргументации", "Слабая работа с возражениями"],
            "recommendations": ["Изучить техники работы с возражениями", "Подготовить больше аргументов"],
            "growth_areas": ["Психология продаж", "Техники убеждения"],
            "techniques_used": ["Вопросы", "Презентация"],
            "missed_opportunities": ["Не использовал социальные доказательства"],
            "overall_rating": "6"
        }

def create_docx_report(report_data: Dict[str, Any], situation: Dict[str, str], conversation_history: List[Dict[str, str]]) -> bytes:
    """Создает отчет в формате DOCX"""
    doc = Document()
    
    # Заголовок
    title = doc.add_heading('Отчет по тренировке переговоров', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Информация о сессии
    doc.add_heading('Информация о сессии', level=1)
    doc.add_paragraph(f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    doc.add_paragraph(f"Ситуация: {situation['situation']}")
    doc.add_paragraph(f"Продукт: {situation['product']}")
    doc.add_paragraph(f"Цель: {situation['manager_goal']}")
    
    # Резюме
    doc.add_heading('Резюме переговоров', level=1)
    doc.add_paragraph(report_data.get('summary', 'Резюме не доступно'))
    
    # Сильные стороны
    doc.add_heading('Сильные стороны', level=1)
    strengths = report_data.get('strengths', [])
    for strength in strengths:
        doc.add_paragraph(f"• {strength}", style='List Bullet')
    
    # Области для улучшения
    doc.add_heading('Области для улучшения', level=1)
    weaknesses = report_data.get('weaknesses', [])
    for weakness in weaknesses:
        doc.add_paragraph(f"• {weakness}", style='List Bullet')
    
    # Рекомендации
    doc.add_heading('Рекомендации', level=1)
    recommendations = report_data.get('recommendations', [])
    for rec in recommendations:
        doc.add_paragraph(f"• {rec}", style='List Bullet')
    
    # Области роста
    doc.add_heading('Области для развития', level=1)
    growth_areas = report_data.get('growth_areas', [])
    for area in growth_areas:
        doc.add_paragraph(f"• {area}", style='List Bullet')
    
    # Использованные техники
    doc.add_heading('Использованные техники', level=1)
    techniques = report_data.get('techniques_used', [])
    for technique in techniques:
        doc.add_paragraph(f"• {technique}", style='List Bullet')
    
    # Упущенные возможности
    doc.add_heading('Упущенные возможности', level=1)
    missed = report_data.get('missed_opportunities', [])
    for opportunity in missed:
        doc.add_paragraph(f"• {opportunity}", style='List Bullet')
    
    # Общая оценка
    doc.add_heading('Общая оценка', level=1)
    rating = report_data.get('overall_rating', 'Не оценено')
    doc.add_paragraph(f"Оценка: {rating}/10")
    
    # История диалога
    doc.add_heading('История диалога', level=1)
    for i, msg in enumerate(conversation_history, 1):
        role = "Менеджер" if msg['role'] == 'manager' else "Клиент"
        doc.add_paragraph(f"{i}. {role}: {msg['content']}")
    
    # Сохраняем в байты
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    return docx_bytes.getvalue()

# Streamlit интерфейс
def main():
    st.set_page_config(
        page_title="Тренажер переговоров",
        page_icon="💼",
        layout="wide"
    )
    
    st.title("💼 Тренажер переговоров с ИИ")
    st.markdown("---")
    
    # Инициализация сессии
    if 'conversation_history' not in st.session_state:
        st.session_state.conversation_history = []
    if 'situation' not in st.session_state:
        st.session_state.situation = None
    if 'dialogue_agent' not in st.session_state:
        st.session_state.dialogue_agent = None
    if 'end_detector' not in st.session_state:
        st.session_state.end_detector = None
    if 'report_generator' not in st.session_state:
        st.session_state.report_generator = None
    if 'dialogue_ended' not in st.session_state:
        st.session_state.dialogue_ended = False
    if 'report_data' not in st.session_state:
        st.session_state.report_data = None
    if 'message_key' not in st.session_state:
        st.session_state.message_key = 0
    if 'user_message' not in st.session_state:
        st.session_state.user_message = ""
    
    # Боковая панель с информацией
    with st.sidebar:
        st.header("📋 Информация о сессии")
        
        if st.session_state.situation:
            st.subheader("Текущая ситуация")
            st.write(f"**Роль:** {st.session_state.situation['manager_role']}")
            st.write(f"**Продукт:** {st.session_state.situation['product']}")
            st.write(f"**Цель:** {st.session_state.situation['manager_goal']}")
            st.write(f"**Возражения клиента:** {st.session_state.situation['client_concerns']}")
        
        st.markdown("---")
        
        if st.button("🔄 Новая тренировка", use_container_width=True):
            st.session_state.conversation_history = []
            st.session_state.situation = None
            st.session_state.dialogue_agent = None
            st.session_state.end_detector = None
            st.session_state.report_generator = None
            st.session_state.dialogue_ended = False
            st.session_state.report_data = None
            st.session_state.message_key = 0
            st.session_state.user_message = ""
            st.rerun()
    
    # Основная область
    col1, col2 = st.columns([2, 1])
    
    with col1:
        if not st.session_state.situation:
            st.info("🎯 Нажмите 'Начать тренировку' чтобы создать новую ситуацию для переговоров")
            
            if st.button("🚀 Начать тренировку", use_container_width=True):
                with st.spinner("Создаю ситуацию для тренировки..."):
                    situation_generator = SituationGenerator(giga)
                    st.session_state.situation = situation_generator.generate_situation()
                    st.session_state.dialogue_agent = DialogueAgent(giga, st.session_state.situation)
                    st.session_state.end_detector = DialogueEndDetector(giga, st.session_state.situation)
                    st.session_state.report_generator = ReportGenerator(giga, st.session_state.situation)
                    st.rerun()
        
        else:
            # Отображение ситуации
            st.subheader("🎭 Ситуация")
            st.info(st.session_state.situation['situation'])
            
            # Отображение диалога
            st.subheader("💬 Диалог")
            
            # Контейнер для диалога
            dialogue_container = st.container()
            
            with dialogue_container:
                for i, msg in enumerate(st.session_state.conversation_history):
                    if msg['role'] == 'manager':
                        st.markdown(f"**👤 Вы (менеджер):** {msg['content']}")
                    else:
                        st.markdown(f"**🤖 Клиент:** {msg['content']}")
                    st.markdown("---")
            
            # Поле ввода сообщения
            if not st.session_state.dialogue_ended:
                user_message = st.text_area(
                    "💭 Ваше сообщение:", 
                    value=st.session_state.user_message,
                    height=100, 
                    placeholder="Введите ваше сообщение клиенту...",
                    key=f"message_input_{st.session_state.message_key}"
                )
                # Очищаем поле ввода
                st.session_state.user_message = ""
                
                col_btn1, col_btn2 = st.columns([1, 1])
                
                with col_btn1:
                    if st.button("📤 Отправить", use_container_width=True):
                        if user_message.strip():
                            # Добавляем сообщение менеджера
                            st.session_state.conversation_history.append({
                                'role': 'manager',
                                'content': user_message.strip()
                            })
                            
                            # Получаем ответ клиента
                            with st.spinner("Клиент думает..."):
                                client_response = st.session_state.dialogue_agent.respond_as_client(
                                    st.session_state.conversation_history
                                )
                                
                                st.session_state.conversation_history.append({
                                    'role': 'client',
                                    'content': client_response
                                })
                            
                            # Проверяем, пора ли завершить диалог
                            if st.session_state.end_detector.should_end_dialogue(st.session_state.conversation_history):
                                st.session_state.dialogue_ended = True
                                with st.spinner("Анализирую результаты..."):
                                    st.session_state.report_data = st.session_state.report_generator.generate_report(
                                        st.session_state.conversation_history
                                    )
                            
                            st.session_state.message_key += 1
                            st.rerun()
                
                with col_btn2:
                    if st.button("🏁 Завершить диалог", use_container_width=True):
                        st.session_state.dialogue_ended = True
                        with st.spinner("Анализирую результаты..."):
                            st.session_state.report_data = st.session_state.report_generator.generate_report(
                                st.session_state.conversation_history
                            )
                        st.session_state.message_key += 1
                        st.rerun()
            else:
                st.success("✅ Диалог завершен! Сгенерируйте отчет для анализа результатов.")
    
    with col2:
        st.subheader("📊 Статистика")
        
        if st.session_state.conversation_history:
            manager_messages = len([msg for msg in st.session_state.conversation_history if msg['role'] == 'manager'])
            client_messages = len([msg for msg in st.session_state.conversation_history if msg['role'] == 'client'])
            
            st.metric("Сообщений менеджера", manager_messages)
            st.metric("Сообщений клиента", client_messages)
            st.metric("Всего сообщений", len(st.session_state.conversation_history))
        
        if st.session_state.dialogue_ended and st.session_state.report_data:
            st.subheader("📈 Быстрый анализ")
            
            rating = st.session_state.report_data.get('overall_rating', 'N/A')
            st.metric("Общая оценка", f"{rating}/10")
            
            strengths_count = len(st.session_state.report_data.get('strengths', []))
            weaknesses_count = len(st.session_state.report_data.get('weaknesses', []))
            
            st.metric("Сильные стороны", strengths_count)
            st.metric("Области роста", weaknesses_count)
            
            # Кнопка для генерации отчета
            if st.button("📄 Скачать отчет", use_container_width=True):
                docx_bytes = create_docx_report(
                    st.session_state.report_data,
                    st.session_state.situation,
                    st.session_state.conversation_history
                )
                
                st.download_button(
                    label="💾 Скачать DOCX отчет",
                    data=docx_bytes,
                    file_name=f"отчет_переговоры_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )

if __name__ == "__main__":
    main() 